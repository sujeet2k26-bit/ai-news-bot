"""
generate_prd.py
───────────────
Generates the Telegram Bot Factory PRD as a Word document.
Run once: python generate_prd.py
Output:   docs/Telegram_Bot_Factory_PRD.docx
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

os.makedirs("docs", exist_ok=True)

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(2.8)
section.right_margin  = Cm(2.8)


# ── Style helpers ─────────────────────────────────────────────────────────────

def set_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    p.runs[0].font.name = "Calibri"
    if level == 1:
        p.runs[0].font.size  = Pt(18)
        p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    elif level == 2:
        p.runs[0].font.size  = Pt(14)
        p.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    elif level == 3:
        p.runs[0].font.size  = Pt(12)
        p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p


def body(text, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
    run = p.add_run(text)
    set_font(run, size=10.5)
    return p


def bullet(text, level=1):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.5 * level)
    run = p.add_run(text)
    set_font(run, size=10.5)
    return p


def note(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    run = p.add_run("📌  " + text)
    set_font(run, size=10, italic=True, color=(0x70, 0x70, 0x70))
    return p


def code_block(lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    for line in lines:
        run = p.add_run(line + "\n")
        set_font(run, name="Courier New", size=9, color=(0x1A, 0x1A, 0x2E))
    p.paragraph_format.space_after = Pt(4)


def divider():
    p = doc.add_paragraph("─" * 80)
    p.runs[0].font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    p.runs[0].font.size = Pt(8)


def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Blue header bg
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "2E74B5")
        tcPr.append(shd)
    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        fill = "F2F7FB" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), fill)
            tcPr.append(shd)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


# ═══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_p.add_run("Telegram Bot Factory")
set_font(r, size=32, bold=True, color=(0x1F, 0x49, 0x7D))

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub_p.add_run("Product Requirements Document")
set_font(r2, size=16, color=(0x2E, 0x74, 0xB5))

doc.add_paragraph()
meta_p = doc.add_paragraph()
meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = meta_p.add_run("Version 1.0  ·  April 2026  ·  Confidential")
set_font(r3, size=11, italic=True, color=(0x70, 0x70, 0x70))

doc.add_paragraph()
doc.add_paragraph()

tagline_p = doc.add_paragraph()
tagline_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = tagline_p.add_run(
    "A multi-bot publishing platform that automates content aggregation,\n"
    "AI-powered generation, human review, and Telegram channel publishing\n"
    "across any domain — from a single shared backend."
)
set_font(r4, size=12, italic=True, color=(0x40, 0x40, 0x40))

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ═══════════════════════════════════════════════════════════════════════════════

heading("Table of Contents", level=1)
toc_items = [
    ("1.", "Executive Summary"),
    ("2.", "Product Vision & Goals"),
    ("3.", "System Architecture"),
    ("4.", "Content Pipeline — Detailed Flow"),
    ("5.", "Bot Specifications"),
    ("6.", "Technical Stack"),
    ("7.", "API Integrations"),
    ("8.", "Reviewer Interface"),
    ("9.", "Hosting Options & Cost Analysis"),
    ("10.", "Scalability Roadmap"),
    ("11.", "Security & Guardrails"),
    ("12.", "Development Roadmap"),
    ("13.", "Operational Runbook"),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    r_num = p.add_run(f"  {num}  ")
    set_font(r_num, bold=True, size=10.5, color=(0x2E, 0x74, 0xB5))
    r_title = p.add_run(title)
    set_font(r_title, size=10.5)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════════════════════════════

heading("1. Executive Summary")
body(
    "Telegram Bot Factory is a production-grade, multi-bot content publishing platform "
    "built on Python. It automates the full lifecycle of digital content — from ingestion "
    "of raw news feeds through AI-powered post generation, mandatory human review, and "
    "final publication to Telegram channels."
)
body(
    "The platform currently operates three live bots — an AI & Technology news bot, a "
    "Bollywood entertainment bot, and a daily Hindi astrology (panchang) bot — all driven "
    "by a single shared backend. The architecture is deliberately bot-agnostic: adding a "
    "new domain bot requires only one configuration entry and one prompt file, with zero "
    "changes to the core pipeline."
)
body(
    "Every post passes through a mandatory human review gate before publishing. No content "
    "goes live without explicit reviewer approval via Telegram inline buttons, eliminating "
    "the risk of AI-generated errors or policy violations reaching an audience."
)

heading("Key Metrics (Current)", level=2)
add_table(
    ["Metric", "Value"],
    [
        ["Active Bots", "3 (AI News, Bollywood, Astrology)"],
        ["Total Telegram Channels", "3"],
        ["Posts per Day", "4 (1 AI + 2 Bollywood + 1 Astrology)"],
        ["AI Model", "Gemini 2.5 Pro via Euri API"],
        ["Free Tier Token Budget", "200,000 tokens/day"],
        ["Est. tokens per digest post", "~2,000 tokens"],
        ["Est. tokens per astrology post", "~1,600 tokens"],
        ["Human Review Required", "Yes — 100% of posts"],
        ["Database", "SQLite (local)"],
        ["Hosting", "Local Windows machine"],
    ],
    col_widths=[6, 10],
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 2. PRODUCT VISION & GOALS
# ═══════════════════════════════════════════════════════════════════════════════

heading("2. Product Vision & Goals")

heading("Vision", level=2)
body(
    "Build the fastest, most reliable way to launch and operate a high-quality AI-powered "
    "Telegram content channel in any domain or language — with human oversight built in "
    "from day one."
)

heading("Core Goals", level=2)
bullet("Bot-agnostic pipeline: one codebase powers any number of bots across any domain")
bullet("Mandatory human review: no post ever published without explicit approval")
bullet("Quality over quantity: virality scoring, source diversity caps, and guardrails ensure only the best content is surfaced")
bullet("30-minute bot onboarding: adding a new domain bot should take under 30 minutes")
bullet("Cost-efficient: Euri free tier (200K tokens/day) covers current post volume with room to scale")

heading("Non-Goals (Current Phase)", level=2)
bullet("No subscriber management or per-user personalisation")
bullet("No real-time trending or social media signal tracking (Phase 2)")
bullet("No web dashboard — review is Telegram-only")
bullet("No multi-language support within a single bot")

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 3. SYSTEM ARCHITECTURE
# ═══════════════════════════════════════════════════════════════════════════════

heading("3. System Architecture")

heading("3.1 High-Level Architecture Diagram", level=2)
body("The platform follows a layered pipeline architecture:")
doc.add_paragraph()

code_block([
    "┌─────────────────────────────────────────────────────────────────────────┐",
    "│                        TELEGRAM BOT FACTORY                            │",
    "│                         (Single Python Process)                         │",
    "├─────────────────────────┬───────────────────────────────────────────────┤",
    "│    INGESTION LAYER      │              PROCESSING LAYER                 │",
    "│                         │                                               │",
    "│  ┌──────────────────┐   │   ┌────────────┐    ┌───────────────────┐    │",
    "│  │  RSS Fetcher     │──▶│──▶│ Guardrails │───▶│  Virality Scorer  │    │",
    "│  │  (30+ sources)   │   │   │ (6 checks) │    │  + Diversity Caps │    │",
    "│  └──────────────────┘   │   └────────────┘    └────────┬──────────┘    │",
    "│  ┌──────────────────┐   │                              │               │",
    "│  │ Panchang Scraper │──▶│──────────────────────────────┘               │",
    "│  │ (Drik Panchang)  │   │                                               │",
    "│  └──────────────────┘   │              GENERATION LAYER                │",
    "│                         │                                               │",
    "│  ┌──────────────────┐   │   ┌──────────────────────────────────────┐   │",
    "│  │  Deduplicator    │   │   │  Gemini 2.5 Pro (via Euri API)       │   │",
    "│  │  (hash-based)    │   │   │  - Digest text generation            │   │",
    "│  └──────────────────┘   │   │  - Cover image generation            │   │",
    "│                         │   │  - Alt headline (A/B test)           │   │",
    "│                         │   │  - Edit instruction apply            │   │",
    "│                         │   └──────────────────────────────────────┘   │",
    "├─────────────────────────┴───────────────────────────────────────────────┤",
    "│                         REVIEW & PUBLISH LAYER                          │",
    "│                                                                          │",
    "│  ┌──────────────────────────────────────────────────────────────────┐   │",
    "│  │                  Human Review Interface (Telegram)               │   │",
    "│  │  Reviewer Chat ──▶ [✅ Approve] [❌ Reject] [📝 Edit] [/generate]│   │",
    "│  └───────────────────────────┬──────────────────────────────────────┘   │",
    "│                              │  Approved                                 │",
    "│                              ▼                                           │",
    "│  ┌────────────────┐  ┌───────────────┐  ┌────────────────────────────┐  │",
    "│  │ @ai26news      │  │@bollywood_    │  │ @astrochhayah              │  │",
    "│  │ (AI News Bot)  │  │daily_gossip   │  │ (Astrology Bot)            │  │",
    "│  └────────────────┘  └───────────────┘  └────────────────────────────┘  │",
    "└──────────────────────────────────────────────────────────────────────────┘",
    "",
    "External Services:  Euri API (Gemini) · Telegram Bot API · Drik Panchang",
    "Storage:            SQLite DB · logs/ directory",
    "Scheduler:          APScheduler (IST timezone) — 6 jobs/day",
])

heading("3.2 Component Breakdown", level=2)
add_table(
    ["Component", "Module", "Responsibility"],
    [
        ["RSS Fetcher", "aggregator/rss_fetcher.py", "Polls 30+ RSS feeds every 6 hours, stores raw articles"],
        ["Panchang Scraper", "aggregator/panchang_fetcher.py", "Scrapes Drik Panchang for daily tithi/nakshatra with cookie auth"],
        ["Deduplicator", "aggregator/dedup.py", "Hash-based dedup — skips articles seen in last 7 days"],
        ["Guardrails", "guardrails/content_filter.py", "Pre + post generation checks across 6 safety categories"],
        ["Virality Scorer", "scoring/virality.py", "Scores by recency + overlap + keywords × source weight"],
        ["Fallback Selector", "scoring/fallback.py", "Best-available articles when no article hits score ≥ 60"],
        ["Generator", "generator/claude_client.py", "Calls Euri API for text + image + alt headline + edit"],
        ["Image Card", "generator/image_card.py", "Pillow-based social card generator (auto-height, Devanagari)"],
        ["Publisher", "publisher/telegram_bot.py", "Bot-agnostic HTML publisher + Markdown→HTML converter"],
        ["Review Interface", "publisher/review_interface.py", "All reviewer commands, Approve/Reject flow, /killstale"],
        ["Scheduler", "scheduler/jobs.py", "APScheduler — 6 daily jobs, IST timezone, all active bots"],
        ["Database", "db/models.py + database.py", "SQLite — Article, Post, PublishLog models"],
        ["Bot Registry", "config/bots.json", "Single source of truth — defines every bot's config"],
    ],
    col_widths=[4, 5, 8],
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 4. CONTENT PIPELINE — DETAILED FLOW
# ═══════════════════════════════════════════════════════════════════════════════

heading("4. Content Pipeline — Detailed Flow")

heading("4.1 Standard Pipeline (AI News & Bollywood)", level=2)
code_block([
    "STEP 1: FETCH (every 6 hours via scheduler)",
    "  RSS Fetcher polls all sources in config/sources_<bot>.json",
    "  New articles saved to DB with status = 'new'",
    "  Deduplicator skips articles seen in last 7 days (hash match)",
    "",
    "STEP 2: FILTER (immediately after fetch)",
    "  Guardrail check on each article:",
    "    ✗ Hate speech / communal content        → status = 'blocked', logged",
    "    ✗ Sexual / adult content                → status = 'blocked', logged",
    "    ✗ Graphic violence                      → status = 'blocked', logged",
    "    ✗ Political propaganda / elections      → status = 'blocked', logged",
    "    ✗ Unverified celebrity defamation       → status = 'blocked', logged",
    "    ✗ Clickbait (non-whitelisted source)    → flagged for review",
    "    ✗ Off-topic (AI filter for ai_news bot) → status = 'blocked', logged",
    "    ✓ Passes all checks                     → status = 'scored'",
    "",
    "STEP 3: SCORE",
    "  virality_score = (recency + source_overlap + keyword_weight) × source_multiplier",
    "    recency:        last 6h = 30pts, 12h = 20pts, 24h = 10pts, 48h+ = 0pts",
    "    source_overlap: same story in 3+ sources = 30pts, 2 sources = 15pts",
    "    keyword_weight: 3+ trending keyword matches = 20pts",
    "    multipliers:    ai_official ×1.5 | ai_research ×1.3 | ai_newsletter ×1.2",
    "                    ai_technology ×1.0 | ai_community ×0.7",
    "  Threshold: score ≥ 60 → 'selected' | below threshold → fallback picks best available",
    "",
    "STEP 4: SELECT (diversity caps applied in order)",
    "  Cap 1: No exact duplicate titles",
    "  Cap 2: Max 2 articles per source",
    "  Cap 3: Max 2 articles per movie/topic (prevents box-office flooding)",
    "  → Final selection: top 5 articles",
    "",
    "STEP 5: GENERATE",
    "  Gemini 2.5 Pro generates digest post text (~2,000 tokens)",
    "  Post-generation guardrail check on generated text",
    "  Gemini image model generates cover image (1024×1024)",
    "  Alt headline variant generated (A/B testing, ~80 tokens)",
    "  Read more URLs injected after each 📌 Source line",
    "  Post saved to DB with status = 'pending_review'",
    "",
    "STEP 6: HUMAN REVIEW",
    "  Post sent to reviewer's Telegram chat:",
    "    - Cover image with header caption",
    "    - Full digest text with ✅ Approve / ❌ Reject / 📝 Alt Headline buttons",
    "  Reviewer options:",
    "    ✅ Approve  → immediate publish to channel",
    "    ❌ Reject   → bot asks for reason → logged",
    "    📝 Edit     → type instruction → Gemini rewrites → re-review",
    "    📝 Alt      → swap in B headline → re-review",
    "",
    "STEP 7: PUBLISH",
    "  publish_post() calls Telegram Bot API with ParseMode.HTML",
    "  DB updated: status = 'published', timestamp, channel recorded",
    "  PublishLog entry created",
])

heading("4.2 Astrology Pipeline (Daily Panchang)", level=2)
code_block([
    "STEP 1: FETCH (6:00 AM IST via scheduler)",
    "  panchang_fetcher.py scrapes drikpanchang.com with session cookies",
    "  Extracts: Tithi, Nakshatra, Yoga, Karana, Paksha, Weekday, Festivals",
    "  Fallback: if scraping fails, builds minimal context from today's date",
    "  Article saved to DB",
    "",
    "STEP 2: GENERATE (single article, no scoring needed)",
    "  Gemini 2.5 Pro writes Hinglish panchang post (~1,600 tokens)",
    "  Sections: Meaning, Daily Insight, Remedy, Tip of the Day",
    "  Spiritual cover image generated (tithi-themed, Madhubani/Tanjore style)",
    "",
    "STEP 3: HUMAN REVIEW → same flow as standard pipeline",
    "",
    "STEP 4: PUBLISH + OPTIONAL IMAGE CARD",
    "  Text post published to @astrochhayah",
    "  Reviewer can run /card to generate and publish an Instagram/WhatsApp card",
    "  Card: Pillow-rendered 1080px wide, auto-height, Nirmala UI font (Devanagari)",
])

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 5. BOT SPECIFICATIONS
# ═══════════════════════════════════════════════════════════════════════════════

heading("5. Bot Specifications")

add_table(
    ["Property", "AI News Bot", "Bollywood Buzz Bot", "Daily Astrology Bot"],
    [
        ["Bot ID", "ai_news", "bollywood", "astrology"],
        ["Telegram Channel", "@ai26news", "@bollywood_daily_gossip", "@astrochhayah"],
        ["Language", "English", "Hindi/Hinglish", "Hindi/Hinglish"],
        ["Post Type", "Digest (5 stories)", "Digest (5 stories)", "Single panchang post"],
        ["Posts per Day", "1", "2", "1"],
        ["Schedule", "7:00 AM IST", "7:00 AM + 6:00 PM IST", "6:00 AM IST"],
        ["Data Source", "RSS feeds (30+ sources)", "RSS feeds (17 sources)", "Drik Panchang (scraped)"],
        ["Scoring", "Virality + source weights", "Virality + source weights", "None (single daily item)"],
        ["Guardrails", "All 6 + AI relevance filter", "All 6", "All 6"],
        ["Image Card", "No", "No", "Yes (/card command)"],
        ["A/B Headlines", "Yes", "Yes", "No"],
        ["Reviewer Chat", "chat ID: 543925804", "chat ID: 543925804", "chat ID: 8622681895"],
        ["Status", "Active", "Active", "Built, inactive"],
    ],
    col_widths=[4, 4, 4, 4],
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 6. TECHNICAL STACK
# ═══════════════════════════════════════════════════════════════════════════════

heading("6. Technical Stack")

add_table(
    ["Layer", "Technology", "Version", "Purpose"],
    [
        ["Runtime", "Python", "3.10+", "Core application language"],
        ["Telegram SDK", "python-telegram-bot", "21.10", "Bot polling, handlers, inline buttons"],
        ["AI Text", "Gemini 2.5 Pro", "via Euri API", "Digest generation, edit instructions"],
        ["AI Images", "Gemini Image Model", "via Euri API", "Cover image generation"],
        ["API Gateway", "Euri (euron.one)", "OpenAI-compatible", "200K free tokens/day, 200+ models"],
        ["Scheduling", "APScheduler", "3.10.4", "6 daily jobs, IST timezone"],
        ["Database", "SQLite", "via SQLAlchemy 2.x", "Articles, Posts, PublishLog"],
        ["RSS Parsing", "feedparser", "6.0.11", "All RSS/Atom feed ingestion"],
        ["Web Scraping", "BeautifulSoup4 + lxml", "4.14.3 / 6.0.2", "Drik Panchang scraping"],
        ["Astronomy", "ephem", "4.1.6", "Sun/Moon calculations for panchang fallback"],
        ["Image Rendering", "Pillow", "10.0+", "Astrology image card generation"],
        ["HTTP Client", "requests", "2.31.0", "API calls, image downloads"],
        ["Env Config", "python-dotenv", "1.0.0", "Load .env secrets"],
    ],
    col_widths=[3.5, 4, 3.5, 5.5],
)

heading("6.1 Database Schema", level=2)
code_block([
    "Table: articles",
    "  id              INTEGER PRIMARY KEY",
    "  title           TEXT",
    "  url             TEXT UNIQUE",
    "  source_name     TEXT",
    "  summary         TEXT",
    "  published_at    DATETIME",
    "  bot_id          TEXT              -- which bot this article belongs to",
    "  status          TEXT              -- new | scored | selected | selected_fallback",
    "                                    -- blocked | used",
    "  virality_score  FLOAT",
    "  url_hash        TEXT UNIQUE       -- for deduplication",
    "  created_at      DATETIME",
    "",
    "Table: posts",
    "  id              INTEGER PRIMARY KEY",
    "  bot_id          TEXT",
    "  article_id      INTEGER FK → articles.id",
    "  content         TEXT              -- full AI-generated post (Telegram markdown)",
    "  image_url       TEXT              -- Euri-generated cover image URL",
    "  headline_b      TEXT              -- alt headline for A/B testing",
    "  status          TEXT              -- pending_review | approved | published | rejected",
    "  reviewed_at     DATETIME",
    "  published_at    DATETIME",
    "  reject_reason   TEXT",
    "  created_at      DATETIME",
    "",
    "Table: publish_log",
    "  id              INTEGER PRIMARY KEY",
    "  post_id         INTEGER FK → posts.id",
    "  bot_id          TEXT",
    "  action          TEXT              -- published | rejected | skipped",
    "  notes           TEXT",
    "  timestamp       DATETIME",
])

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 7. API INTEGRATIONS
# ═══════════════════════════════════════════════════════════════════════════════

heading("7. API Integrations")

add_table(
    ["Service", "Usage", "Auth Method", "Limits / Notes"],
    [
        ["Telegram Bot API", "Send/receive messages, publish posts, poll updates", "Bot Token (per bot)", "No hard rate limit; 30 msgs/sec per bot"],
        ["Euri API (Gemini)", "Text generation + image generation", "EURI_API_KEY", "200K tokens/day free; resets at UTC midnight"],
        ["Drik Panchang", "Scrape daily tithi/nakshatra/yoga data", "Session cookies (annual)", "Cookies expire ~1 year; fallback to ephem calc"],
        ["NewsAPI", "Supplementary article ingestion (Phase 2)", "NEWS_API_KEY", "100 req/day free tier"],
        ["Reddit API", "Virality signals from r/MachineLearning etc. (Phase 2)", "OAuth2 client credentials", "60 req/min free"],
    ],
    col_widths=[3.5, 5, 4, 4],
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 8. REVIEWER INTERFACE
# ═══════════════════════════════════════════════════════════════════════════════

heading("8. Reviewer Interface")

heading("8.1 Review Commands", level=2)
add_table(
    ["Command", "Description", "Works From"],
    [
        ["/generate", "Generate a new post for the current bot", "Any reviewer chat"],
        ["/generate <bot_id>", "Generate for any bot (ai_news / bollywood / astrology)", "Any reviewer chat"],
        ["/card", "Generate + publish Instagram/Facebook image card (astrology)", "Astrology reviewer chat"],
        ["/card full", "Full auto-height card for WhatsApp", "Astrology reviewer chat"],
        ["/bulkcard <days>", "Bulk-generate cards for N days from today (max 90); each date uses its own Drik Panchang scrape. Sends a single alert if any dates fell back to ephem.", "Astrology reviewer chat"],
        ["/edit [post_id]", "Edit post with natural language instruction before approving", "Any reviewer chat"],
        ["/pending", "List all posts awaiting review", "Any reviewer chat"],
        ["/preview <id>", "Show full content of a specific post", "Any reviewer chat"],
        ["/sources <id>", "Show the source article used for a post", "Any reviewer chat"],
        ["/skip <id>", "Skip a post without publishing", "Any reviewer chat"],
        ["/killstale", "Kill stale Python processes (fixes 409 Conflict)", "Any reviewer chat"],
        ["/help", "Show all available commands", "Any reviewer chat"],
    ],
    col_widths=[4.5, 7, 5],
)

heading("8.2 Review Flow Diagram", level=2)
code_block([
    "Post Generated",
    "      │",
    "      ▼",
    "┌─────────────────────────────────┐",
    "│  Reviewer receives in Telegram: │",
    "│  [Cover Image]                  │",
    "│  Bot: AI News Bot | Post ID: 47 │",
    "│  Source: TechCrunch             │",
    "│  ─────────────────────────────  │",
    "│  [Full digest content]          │",
    "│  [✅ Approve] [❌ Reject]        │",
    "│  [📰 View Source Article]       │",
    "│  [📝 Use Alt Headline (B)]      │",
    "└─────────────────────────────────┘",
    "         │              │",
    "    ✅ Approve      ❌ Reject",
    "         │              │",
    "         ▼              ▼",
    "   Published to    Bot asks:",
    "   @ai26news       'Type reject reason'",
    "   immediately          │",
    "                        ▼",
    "                   Reviewer types reason",
    "                        │",
    "                        ▼",
    "                   Post marked 'rejected'",
    "                   Reason logged to DB",
    "                   + publish_history.log",
])

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 9. HOSTING OPTIONS & COST ANALYSIS
# ═══════════════════════════════════════════════════════════════════════════════

heading("9. Hosting Options & Cost Analysis")

heading("9.1 Resource Requirements", level=2)
body(
    "The application is lightweight. All three bots run as daemon threads inside a "
    "single Python process. Resource usage profile:"
)
add_table(
    ["Resource", "Requirement", "Notes"],
    [
        ["CPU", "< 5% average", "Mostly idle; spikes to ~20% during Gemini API calls (30–60s, 4×/day)"],
        ["RAM", "256–512 MB", "3 bot threads + APScheduler + SQLAlchemy session pool"],
        ["Disk", "< 500 MB", "SQLite DB grows ~1 MB/month; logs ~10 MB/month; card images ~2 MB/day"],
        ["Network", "Minimal outbound", "RSS polling (6h intervals), Telegram API, Euri API calls"],
        ["Uptime", "Near 24/7", "Must be running at 6–8 AM IST and 6–7 PM IST for scheduled posts"],
        ["OS", "Linux preferred", "Windows works locally; Linux is standard for cloud hosting"],
    ],
    col_widths=[3, 4, 9.5],
)

heading("9.2 Hosting Platform Comparison", level=2)
add_table(
    ["Platform", "Plan", "Monthly Cost (USD)", "RAM", "vCPU", "Pros", "Cons"],
    [
        ["DigitalOcean Droplet", "Basic 1GB", "$6", "1 GB", "1", "Simple, reliable, SSH access, $200 free credit for new users", "Manual setup required"],
        ["DigitalOcean Droplet", "Basic 2GB", "$12", "2 GB", "1", "More headroom for future bots", "Slight overkill for current load"],
        ["AWS EC2", "t3.micro", "~$8–10", "1 GB", "2 vCPU", "Free tier 12 months, excellent ecosystem", "More complex setup; costs rise quickly"],
        ["Railway", "Hobby Plan", "$5", "512 MB", "Shared", "Git push to deploy, zero DevOps", "512 MB RAM is tight; cold starts"],
        ["Render", "Starter", "$7", "512 MB", "0.5 vCPU", "Auto-deploy from GitHub, free SSL", "Spins down on inactivity (free tier)"],
        ["Fly.io", "Pay-as-you-go", "~$3–6", "256 MB", "Shared", "Very cheap, global edge", "More complex config; limited free tier"],
        ["Hetzner Cloud", "CX11", "€3.29 (~$3.50)", "2 GB", "1 vCPU", "Best price/performance in EU", "EU datacentres only"],
        ["Local Machine", "— (current)", "$0", "Unlimited", "—", "No cost, easy debug", "No uptime guarantee; tied to machine"],
    ],
    col_widths=[3.5, 2.5, 2.5, 1.5, 1.5, 4, 3],
)

heading("9.3 Recommended Setup by Stage", level=2)

heading("Stage 1 — Development / Testing (Current)", level=3)
bullet("Local Windows machine — $0/month")
bullet("Run python main.py manually; stop with Ctrl+C")
bullet("Suitable for building and testing new bots before going live")

heading("Stage 2 — Production (Single Server)", level=3)
bullet("DigitalOcean Basic Droplet 1GB — $6/month")
bullet("Ubuntu 22.04 LTS, Python 3.11, systemd service for auto-restart")
bullet("SQLite database (local file on droplet)")
bullet("Total monthly cost: $6 + $0 Euri (free tier) = $6/month")
note("This handles up to ~10 bots posting 2–3 times/day comfortably within free Euri quota.")

heading("Stage 3 — Scale (Multiple Servers / Managed DB)", level=3)
bullet("DigitalOcean Droplet 2GB ($12) + Managed PostgreSQL ($15) = $27/month")
bullet("Or AWS t3.small ($16) + RDS PostgreSQL free tier = ~$16/month first year")
bullet("PostgreSQL needed when: multiple server instances, >50 bots, or audit requirements")
bullet("Add a load balancer ($12/month on DigitalOcean) only if running multiple app servers")

heading("9.4 Full Cost Breakdown — Production Scenario", level=2)
add_table(
    ["Item", "Stage 1 (Local)", "Stage 2 (Single VPS)", "Stage 3 (Scale)"],
    [
        ["Server / Compute", "$0", "$6/month", "$12–16/month"],
        ["Database", "$0 (SQLite)", "$0 (SQLite on VPS)", "$15/month (Managed PG)"],
        ["Euri API (AI text + image)", "$0 (free tier)", "$0 (free tier)", "$0–20/month (if >200K tokens/day)"],
        ["Telegram Bot API", "$0 (always free)", "$0", "$0"],
        ["Domain (optional)", "$0", "$10–12/year", "$10–12/year"],
        ["SSL Certificate", "N/A", "$0 (Let's Encrypt)", "$0"],
        ["Monitoring (optional)", "$0", "$0–4/month", "$4–8/month"],
        ["TOTAL", "$0/month", "~$6/month", "~$27–44/month"],
    ],
    col_widths=[5, 3.5, 3.5, 4.5],
)

note(
    "Euri free tier (200K tokens/day) comfortably covers current usage: "
    "4 posts × ~2,000 tokens = ~8,000 tokens/day. You have 25× headroom before "
    "needing a paid plan. At $0.15/1M tokens (Euri paid), 1M tokens/day would cost ~$4.50/day."
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 10. SCALABILITY ROADMAP
# ═══════════════════════════════════════════════════════════════════════════════

heading("10. Scalability Roadmap")

heading("10.1 Adding New Bots", level=2)
body(
    "The platform is designed for zero-friction bot onboarding. The steps to add any "
    "new domain bot:"
)
add_table(
    ["Step", "Action", "File(s)", "Time Estimate"],
    [
        ["1", "Add config entry", "config/bots.json", "5 min"],
        ["2", "Create sources file", "config/sources_<id>.json", "10 min"],
        ["3", "Create prompt file", "generator/prompts_<id>.py", "10 min"],
        ["4", "Add env vars", ".env", "2 min"],
        ["5", "Register token in code", "review_interface.py, telegram_bot.py", "3 min"],
        ["6", "Set active=true, restart", "config/bots.json", "1 min"],
    ],
    col_widths=[1, 4.5, 5.5, 3.5],
)
note("Target: under 30 minutes for any new domain bot.")

heading("10.2 Scaling Triggers", level=2)
add_table(
    ["Trigger", "Symptom", "Solution"],
    [
        ["5+ active bots", "Single process feels crowded", "Split into per-domain Python processes or use asyncio task groups"],
        ["Euri free quota exceeded", "429 rate limit errors in generation", "Upgrade to Euri paid plan (~$0.15/1M tokens) or add queue with rate limiter"],
        ["SQLite write contention", "DB locked errors under concurrent writes", "Migrate to PostgreSQL (schema is SQLAlchemy-compatible — change DATABASE_URL only)"],
        ["Multiple reviewers needed", "One person can't review all bots", "Each bot already supports its own TELEGRAM_<BOT>_REVIEWER_CHAT_ID"],
        ["High traffic / analytics", "Need post engagement data", "Add webhook + analytics DB table in Phase 3"],
    ],
    col_widths=[4, 4.5, 8],
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 11. SECURITY & GUARDRAILS
# ═══════════════════════════════════════════════════════════════════════════════

heading("11. Security & Guardrails")

heading("11.1 Content Guardrails (6 Categories)", level=2)
add_table(
    ["Category", "What is Blocked", "Action"],
    [
        ["Hate Speech / Communal", "Content targeting religions, castes, or ethnic groups", "Hard block + log"],
        ["Sexual / Adult Content", "Explicit or suggestive content", "Hard block + log"],
        ["Graphic Violence", "Gore, brutality descriptions beyond factual reporting", "Hard block + log"],
        ["Political Content", "Elections, propaganda, political targeting", "Hard block + log"],
        ["Unverified Defamation", "Celebrity accusations without named credible source", "Hard block + log"],
        ["Clickbait / Fake News", "Non-whitelisted sources; SHOCKING/EXPOSED headlines", "Flag for human review"],
        ["AI Relevance (ai_news only)", "Community-source articles with no AI keywords", "Skip + log as off_topic"],
    ],
    col_widths=[4.5, 7, 5],
)
body(
    "Two check rounds: pre-generation (saves API tokens) and post-generation "
    "(catches AI hallucinations or accidental policy violations in output)."
)

heading("11.2 Secrets Management", level=2)
bullet(".env file is gitignored — API keys never committed to version control")
bullet("Each bot uses its own Telegram token — compromise of one token doesn't affect others")
bullet("Drik Panchang session cookies stored in .env — expire ~1 year; bot sends alert when renewal is needed")
bullet("No user data stored — the platform only processes public news content")

heading("11.3 Human Review as a Security Layer", level=2)
bullet("100% of posts require explicit human approval before publishing")
bullet("Reviewer sees full content + source article before deciding")
bullet("Reject reason is logged — creates an audit trail for content decisions")
bullet("No auto-publish path exists in the codebase — the approval gate cannot be bypassed")

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 12. DEVELOPMENT ROADMAP
# ═══════════════════════════════════════════════════════════════════════════════

heading("12. Development Roadmap")

add_table(
    ["Phase", "Feature", "Status", "Notes"],
    [
        ["Phase 1", "RSS fetcher (AI + Bollywood)", "✅ Done", ""],
        ["Phase 1", "Hash-based deduplication", "✅ Done", ""],
        ["Phase 1", "6-category guardrail filter", "✅ Done", "Pre + post generation"],
        ["Phase 1", "Virality scoring + source weights", "✅ Done", ""],
        ["Phase 1", "Source + topic diversity caps", "✅ Done", "Max 2 per source/topic"],
        ["Phase 1", "Digest generation via Gemini 2.5 Pro", "✅ Done", "Top 5 articles"],
        ["Phase 1", "Cover image generation", "✅ Done", "Gemini image model"],
        ["Phase 1", "Read more links injection", "✅ Done", "Post-generation URL inject"],
        ["Phase 1", "Human review via Telegram", "✅ Done", "Approve/Reject buttons"],
        ["Phase 1", "On-demand /generate command", "✅ Done", "Cross-bot from any reviewer chat"],
        ["Phase 1", "Publish scheduler (IST)", "✅ Done", "APScheduler, 6 jobs/day"],
        ["Phase 1", "SQLite database", "✅ Done", ""],
        ["Phase 1", "AI News Bot + Bollywood Bot", "✅ Done", "Both active"],
        ["Phase 1", "Daily Astrology Bot (panchang)", "✅ Done", "Built, inactive"],
        ["Phase 1", "Per-bot reviewer chat IDs", "✅ Done", ""],
        ["Phase 1", "All 3 review bots auto-start", "✅ Done", "Daemon threads in main.py"],
        ["Phase 1", "A/B headline testing", "✅ Done", "Alt headline + swap button"],
        ["Phase 1", "/edit command (AI rewrites)", "✅ Done", "Natural language instruction"],
        ["Phase 1", "Astrology image card (/card)", "✅ Done", "Auto-height, Devanagari font"],
        ["Phase 1", "/killstale command", "✅ Done", "PowerShell PID-safe process kill"],
        ["Phase 1", "/bulkcard <days> command", "✅ Done", "Bulk card generation for N days; per-date Drik Panchang scrape; end-of-run ephem alert"],
        ["Phase 2", "Reddit virality signals", "⬜ Planned", "r/MachineLearning, r/Bollywood"],
        ["Phase 2", "NewsAPI integration", "⬜ Planned", "100 req/day free"],
        ["Phase 2", "Subscriber management", "⬜ Planned", "Subscribe/unsubscribe per bot"],
        ["Phase 2", "Web scraping for non-RSS sources", "⬜ Planned", ""],
        ["Phase 2", "Keyword trending tracker", "⬜ Planned", ""],
        ["Phase 2", "Activate Astrology Bot on schedule", "⬜ Planned", "Set active=true in bots.json"],
        ["Phase 3", "WhatsApp Business API integration", "⬜ Future", ""],
        ["Phase 3", "Web dashboard for review", "⬜ Future", "Replace Telegram reviewer"],
        ["Phase 3", "Engagement analytics per post", "⬜ Future", ""],
        ["Phase 3", "Cloud hosting migration", "⬜ Future", "DigitalOcean or AWS"],
        ["Phase 3", "PostgreSQL migration", "⬜ Future", "Change DATABASE_URL only"],
    ],
    col_widths=[2.5, 5.5, 2.5, 6],
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 13. OPERATIONAL RUNBOOK
# ═══════════════════════════════════════════════════════════════════════════════

heading("13. Operational Runbook")

heading("13.1 Daily Schedule (IST)", level=2)
add_table(
    ["Time (IST)", "Action", "Bots"],
    [
        ["6:00 AM", "Fetch panchang + generate astrology post → send to reviewer", "Astrology"],
        ["6:00 AM", "Fetch + score articles from RSS", "AI News, Bollywood"],
        ["7:00 AM", "Generate digest + send to reviewer", "AI News, Bollywood"],
        ["8:00 AM", "Publish approved posts; hold if no response", "AI News, Bollywood"],
        ["12:00 PM", "Midday fetch + retry unapproved morning posts", "Bollywood"],
        ["6:00 PM", "Generate second digest + send to reviewer", "Bollywood"],
        ["7:00 PM", "Publish approved evening post", "Bollywood"],
    ],
    col_widths=[3, 9.5, 4],
)

heading("13.2 Starting and Stopping", level=2)
code_block([
    "# Start all bots (production)",
    "python main.py",
    "",
    "# Start a single bot for testing",
    "python publisher/test_review_interface.py astrology",
    "",
    "# Fetch fresh articles manually",
    "python aggregator/test_fetch.py ai_news",
    "python aggregator/test_fetch.py bollywood",
    "",
    "# Stop everything (Windows)",
    "Ctrl+C   (if running in foreground)",
    "taskkill //F //IM python.exe   (if running in background)",
])

heading("13.3 Troubleshooting", level=2)
add_table(
    ["Problem", "Cause", "Fix"],
    [
        ["409 Conflict error", "Stale Python process still polling same token", "Send /killstale from any reviewer chat, or taskkill //F //IM python.exe"],
        ["403 Forbidden on review send", "Reviewer hasn't sent /start to that bot", "Open Telegram, find the bot, send /start from the reviewer account"],
        ["Chat not found", "TELEGRAM_REVIEWER_CHAT_ID is @username not numeric ID", "Get numeric ID from @userinfobot; update .env"],
        ["Text-only review (no image)", "Euri image URL expired (~5 min TTL)", "Normal fallback — post can still be approved and published"],
        ["Generation fails / empty response", "Euri API key missing or quota exhausted", "Check .env; run python test_euri_connection.py; wait for quota reset at UTC midnight"],
        ["No articles in digest", "No fetched articles or all blocked by guardrails", "Run test_fetch.py; check logs/guardrail_violations.log"],
        ["Panchang scraping fails", "Drik Panchang cookies expired", "Re-copy session cookies from browser → DevTools → Application → Cookies → drikpanchang.com"],
        ["Post content looks wrong", "AI preamble or formatting issue", "Use /edit with a correction instruction; or reject and /generate again"],
        ["/card shows truncated text", "Old cached image_card.py loaded", "importlib.reload() is called automatically in /card handler — should not occur"],
    ],
    col_widths=[4, 5, 7.5],
)

heading("13.4 Deploying to a Cloud VPS", level=2)
code_block([
    "# 1. Provision server (Ubuntu 22.04 on DigitalOcean/AWS/Hetzner)",
    "",
    "# 2. Install Python 3.11",
    "sudo apt update && sudo apt install python3.11 python3.11-venv python3-pip -y",
    "",
    "# 3. Clone the repo",
    "git clone https://github.com/sujeet2k26-bit/telegram-bot-factory.git",
    "cd telegram-bot-factory",
    "",
    "# 4. Create virtual environment and install dependencies",
    "python3.11 -m venv venv",
    "source venv/bin/activate",
    "pip install -r requirements.txt",
    "",
    "# 5. Create .env from template and fill in your keys",
    "cp .env.example .env",
    "nano .env",
    "",
    "# 6. Set up as a systemd service (auto-restart on crash/reboot)",
    "sudo nano /etc/systemd/system/telegram-bot-factory.service",
    "",
    "  [Unit]",
    "  Description=Telegram Bot Factory",
    "  After=network.target",
    "",
    "  [Service]",
    "  User=ubuntu",
    "  WorkingDirectory=/home/ubuntu/telegram-bot-factory",
    "  ExecStart=/home/ubuntu/telegram-bot-factory/venv/bin/python main.py",
    "  Restart=always",
    "  RestartSec=10",
    "",
    "  [Install]",
    "  WantedBy=multi-user.target",
    "",
    "sudo systemctl enable telegram-bot-factory",
    "sudo systemctl start telegram-bot-factory",
    "sudo systemctl status telegram-bot-factory",
    "",
    "# 7. View live logs",
    "journalctl -u telegram-bot-factory -f",
    "# or",
    "tail -f logs/app.log",
])

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# APPENDIX: ENVIRONMENT VARIABLES
# ═══════════════════════════════════════════════════════════════════════════════

heading("Appendix A: Environment Variables Reference")
add_table(
    ["Variable", "Required", "Description"],
    [
        ["EURI_API_KEY", "Yes", "Euri API key for Gemini text + image generation"],
        ["TELEGRAM_AI_BOT_TOKEN", "Yes", "Bot token for AI News Bot"],
        ["TELEGRAM_AI_CHANNEL_ID", "Yes", "Channel ID for AI News (e.g. @ai26news)"],
        ["TELEGRAM_BOLLYWOOD_BOT_TOKEN", "Yes", "Bot token for Bollywood Buzz Bot"],
        ["TELEGRAM_BOLLYWOOD_CHANNEL_ID", "Yes", "Channel ID for Bollywood"],
        ["TELEGRAM_ASTROLOGY_BOT_TOKEN", "Yes", "Bot token for Daily Astrology Bot"],
        ["TELEGRAM_ASTROLOGY_CHANNEL_ID", "Yes", "Channel ID for Astrology"],
        ["TELEGRAM_REVIEWER_CHAT_ID", "Yes", "Default reviewer numeric chat ID (AI News + Bollywood)"],
        ["TELEGRAM_ASTROLOGY_REVIEWER_CHAT_ID", "Optional", "Astrology reviewer chat ID if different account"],
        ["DRIK_SESSION_ID", "Optional", "Drik Panchang session cookie (for authenticated scraping)"],
        ["DRIK_ACCESS_TOKEN", "Optional", "Drik Panchang access token cookie"],
        ["DRIK_GEONAME_ID", "Optional", "Geonames city ID for panchang location (default: 1261481 = New Delhi)"],
        ["NEWS_API_KEY", "Phase 2", "NewsAPI.org key (100 req/day free)"],
        ["REDDIT_CLIENT_ID", "Phase 2", "Reddit app client ID"],
        ["REDDIT_CLIENT_SECRET", "Phase 2", "Reddit app client secret"],
        ["DATABASE_URL", "Optional", "SQLAlchemy DB URL (default: sqlite:///db/ainews.db)"],
        ["LOG_LEVEL", "Optional", "Logging verbosity (default: INFO; use DEBUG for development)"],
    ],
    col_widths=[5.5, 2.5, 8.5],
)

# ── Save ──────────────────────────────────────────────────────────────────────
output_path = os.path.abspath("docs/Telegram_Bot_Factory_PRD.docx")
doc.save(output_path)
print(f"PRD saved: {output_path}")
